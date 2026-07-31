# Copyright 2026 Ajay Rajan
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.

"""DOCX extraction via python-docx."""

from __future__ import annotations

import io
from typing import Optional

from ..types import ExtractedDocument


def extract_docx(data: bytes, *, source: Optional[str] = None) -> ExtractedDocument:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text]
    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    return ExtractedDocument(text=text, format="docx", source=source)
